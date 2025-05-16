import os
import zipfile
import tempfile
import pythoncom
import win32com.client
import pytesseract
import textwrap
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from fpdf import FPDF
from pdf2image import convert_from_path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
import openpyxl
import shutil

# Set these paths as needed for your system:
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
POPPLER_PATH = r'C:\Program Files\poppler-23.11.0\Library\bin'  # Update if needed

# TXT to PDF
def convert_txt_to_pdf(input_file, output_file):
    with open(input_file, 'r', encoding='utf-8') as file:
        text = file.read()
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, text)
    pdf.output(output_file)

# TXT to DOCX
def convert_txt_to_docx(input_file, output_file):
    with open(input_file, 'r', encoding='utf-8') as file:
        text = file.read()
    doc = Document()
    doc.add_paragraph(text)
    doc.save(output_file)

# TXT to PNG (improved)
def convert_txt_to_image(input_file, output_file):
    try:
        with open(input_file, 'r', encoding='utf-8') as file:
            text = file.read()
        font_path = "cour.ttf" if os.path.exists("cour.ttf") else None
        font_size = 14
        line_spacing = 1.2
        margin = 20
        chars_per_line = 80
        lines = []
        for paragraph in text.split('\n'):
            lines.extend(textwrap.wrap(paragraph, width=chars_per_line, replace_whitespace=False))
        if font_path:
            font = ImageFont.truetype(font_path, font_size)
        else:
            font = ImageFont.load_default()
        max_width = max(font.getsize(line)[0] for line in lines) if lines else 0
        total_height = int(len(lines) * font_size * line_spacing)
        img_width = max_width + 2*margin
        img_height = total_height + 2*margin
        image = Image.new('RGB', (img_width, img_height), color='white')
        draw = ImageDraw.Draw(image)
        y = margin
        for line in lines:
            draw.text((margin, y), line, font=font, fill='black')
            y += int(font_size * line_spacing)
        image.save(output_file)
        return True
    except Exception as e:
        print(f"TXT to PNG Error: {e}")
        return False

# DOCX to TXT
def convert_docx_to_txt(input_file, output_file):
    doc = Document(input_file)
    with open(output_file, 'w', encoding='utf-8') as f:
        for para in doc.paragraphs:
            f.write(para.text + '\n')

# DOCX to PDF
def convert_docx_to_pdf(input_file, output_file):
    doc = Document(input_file)
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for para in doc.paragraphs:
        pdf.multi_cell(0, 10, para.text)
    pdf.output(output_file)

# DOCX to PNG (via PDF intermediate)
def convert_docx_to_image(input_file, output_file):
    try:
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as temp_pdf:
            convert_docx_to_pdf(input_file, temp_pdf.name)
        images = convert_from_path(temp_pdf.name, poppler_path=POPPLER_PATH)
        if images:
            images[0].save(output_file, 'PNG')
            return True
        return False
    except Exception as e:
        print(f"DOCX to PNG Error: {e}")
        return False
    finally:
        if os.path.exists(temp_pdf.name):
            os.remove(temp_pdf.name)

def convert_docx_to_excel(input_file, output_file):
    doc = Document(input_file)
    wb = openpyxl.Workbook()
    ws = wb.active
    for i, para in enumerate(doc.paragraphs, start=1):
        ws.cell(row=i, column=1, value=para.text)
    wb.save(output_file)

# PPTX to PDF
def convert_pptx_to_pdf(input_file, output_file):
    pythoncom.CoInitialize()
    powerpoint = win32com.client.Dispatch("PowerPoint.Application")
    presentation = powerpoint.Presentations.Open(os.path.abspath(input_file), WithWindow=False)
    presentation.SaveAs(os.path.abspath(output_file), 32)
    presentation.Close()
    powerpoint.Quit()
    pythoncom.CoUninitialize()

# PPTX to ZIP of PNGs
def convert_pptx_to_zip(input_file, output_file):
    presentation = None
    powerpoint = None
    try:
        pythoncom.CoInitialize()
        powerpoint = win32com.client.Dispatch("PowerPoint.Application")
        presentation = powerpoint.Presentations.Open(os.path.abspath(input_file), WithWindow=False)
        with tempfile.TemporaryDirectory() as temp_dir:
            for i in range(1, presentation.Slides.Count + 1):
                slide = presentation.Slides(i)
                slide.Export(os.path.join(temp_dir, f"slide_{i}.png"), "PNG")
            if not os.listdir(temp_dir):
                raise RuntimeError("No slides were converted to PNG")
            temp_zip = output_file + ".tmp"
            with zipfile.ZipFile(temp_zip, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                for img_name in os.listdir(temp_dir):
                    img_path = os.path.join(temp_dir, img_name)
                    zip_file.write(img_path, img_name)
            if os.path.exists(output_file):
                os.remove(output_file)
            os.rename(temp_zip, output_file)
        return True
    except Exception as e:
        if 'temp_zip' in locals() and os.path.exists(temp_zip):
            os.remove(temp_zip)
        raise RuntimeError(f"PPTX conversion failed: {str(e)}")
    finally:
        try:
            if presentation:
                presentation.Close()
            if powerpoint:
                powerpoint.Quit()
            pythoncom.CoUninitialize()
        except:
            pass

# XLSX to PDF
def convert_excel_to_pdf(input_file, output_file):
    wb = openpyxl.load_workbook(input_file)
    sheet = wb.active
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for row in sheet.iter_rows(values_only=True):
        line = "\t".join(str(cell) for cell in row if cell is not None)
        pdf.multi_cell(0, 10, line)
    pdf.output(output_file)

# XLSX to DOCX
def convert_excel_to_docx(input_file, output_file):
    wb = openpyxl.load_workbook(input_file)
    sheet = wb.active
    doc = Document()
    for row in sheet.iter_rows(values_only=True):
        line = "\t".join(str(cell) for cell in row if cell is not None)
        doc.add_paragraph(line)
    doc.save(output_file)

# XLSX to TXT
def convert_excel_to_txt(input_file, output_file):
    wb = openpyxl.load_workbook(input_file)
    sheet = wb.active
    with open(output_file, 'w', encoding='utf-8') as f:
        for row in sheet.iter_rows(values_only=True):
            line = "\t".join(str(cell) for cell in row if cell is not None)
            f.write(line + '\n')

# XLSX to PNG (table visualization)
def convert_excel_to_image(input_file, output_file):
    try:
        df = pd.read_excel(input_file)
        plt.figure(figsize=(12, min(4 + len(df)*0.3, 20)))
        ax = plt.gca()
        ax.axis('off')
        table = ax.table(cellText=df.values,
                        colLabels=df.columns,
                        cellLoc='center',
                        loc='center')
        table.auto_set_font_size(False)
        table.set_fontsize(8)
        table.scale(1, 1.5)
        plt.savefig(output_file, bbox_inches='tight', dpi=300)
        plt.close()
        return True
    except Exception as e:
        print(f"Excel to PNG Error: {e}")
        return False

# PNG to TXT (OCR)
def convert_png_to_txt(input_file, output_file):
    try:
        img = Image.open(input_file)
        text = pytesseract.image_to_string(img)
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(text)
        return True
    except Exception as e:
        print(f"OCR Error: {e}")
        return False

# PNG/JPG to PDF
def convert_image_to_pdf(input_file, output_file):
    try:
        image = Image.open(input_file)
        if image.mode != 'RGB':
            image = image.convert('RGB')
        pdf = FPDF(unit='pt', format=[image.width, image.height])
        pdf.add_page()
        temp_img = output_file + ".temp.jpg"
        image.save(temp_img)
        pdf.image(temp_img, 0, 0, image.width, image.height)
        pdf.output(output_file)
        os.remove(temp_img)
        return True
    except Exception as e:
        print(f"Image to PDF Error: {e}")
        return False

# PNG to JPG and JPG to PNG
def convert_png_to_jpg(input_file, output_file):
    try:
        image = Image.open(input_file)
        if image.mode != 'RGB':
            image = image.convert('RGB')
        image.save(output_file, 'JPEG', quality=95)
        return True
    except Exception as e:
        print(f"PNG to JPG Error: {e}")
        return False

def convert_jpg_to_png(input_file, output_file):
    try:
        image = Image.open(input_file)
        image.save(output_file, 'PNG')
        return True
    except Exception as e:
        print(f"JPG to PNG Error: {e}")
        return False

# Main Conversion Function
def convert_file(input_file, output_file):
    input_ext = os.path.splitext(input_file)[1].lower()
    output_ext = os.path.splitext(output_file)[1].lower()

    conversion_map = {
        ('.txt', '.pdf'): convert_txt_to_pdf,
        ('.txt', '.docx'): convert_txt_to_docx,
        ('.txt', '.png'): convert_txt_to_image,
        ('.docx', '.txt'): convert_docx_to_txt,
        ('.docx', '.pdf'): convert_docx_to_pdf,
        ('.docx', '.png'): convert_docx_to_image,
        ('.docx', '.xlsx'): convert_docx_to_excel,
        ('.pptx', '.pdf'): convert_pptx_to_pdf,
        ('.pptx', '.zip'): convert_pptx_to_zip,
        ('.xlsx', '.pdf'): convert_excel_to_pdf,
        ('.xlsx', '.docx'): convert_excel_to_docx,
        ('.xlsx', '.txt'): convert_excel_to_txt,
        ('.xlsx', '.png'): convert_excel_to_image,
        ('.png', '.txt'): convert_png_to_txt,
        ('.png', '.pdf'): convert_image_to_pdf,
        ('.png', '.jpg'): convert_png_to_jpg,
        ('.jpg', '.pdf'): convert_image_to_pdf,
        ('.jpg', '.png'): convert_jpg_to_png,
    }

    if input_ext == output_ext:
        raise ValueError("Input and output formats cannot be the same")

    func = conversion_map.get((input_ext, output_ext))
    if not func:
        raise ValueError(f"Unsupported conversion: {input_ext} to {output_ext}")

    try:
        os.makedirs(os.path.dirname(output_file), exist_ok=True)
        func(input_file, output_file)
        if not os.path.exists(output_file):
            raise RuntimeError("Output file was not created")
        if output_ext == '.zip':
            with zipfile.ZipFile(output_file, 'r') as zip_ref:
                if len(zip_ref.namelist()) == 0:
                    os.remove(output_file)
                    raise RuntimeError("Converted ZIP file is empty")
        return True
    except Exception as e:
        if os.path.exists(output_file):
            if os.path.isdir(output_file):
                shutil.rmtree(output_file)
            else:
                os.remove(output_file)
        raise RuntimeError(f"Conversion failed: {e}")
